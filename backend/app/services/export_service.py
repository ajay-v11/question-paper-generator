import io
from typing import Dict, Any, List
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.models.question import (
    GeneratedQuestions,
    MCQuestion,
    FillBlankQuestion,
    ShortQuestion,
    LongQuestion,
)
from app.models.paper import PaperResponse


class ExportService:
    @staticmethod
    def generate_pdf(paper: PaperResponse, subject_name: str) -> bytes:
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)

        styles = getSampleStyleSheet()
        title_style = styles["Heading1"]
        title_style.alignment = 1  # Center alignment

        header_style = styles["Heading2"]
        normal_style = styles["Normal"]

        # Create custom styles
        question_style = ParagraphStyle(
            "Question", parent=styles["Normal"], spaceAfter=6, fontSize=11, leading=14
        )

        option_style = ParagraphStyle(
            "Option", parent=styles["Normal"], leftIndent=20, spaceAfter=2
        )

        elements = []

        elements.append(Paragraph("Question Paper", title_style))
        elements.append(Spacer(1, 12))

        data = [
            [f"Subject: {subject_name}", f"Difficulty: {paper.difficulty.title()}"],
            [
                f"Units: {', '.join(map(str, paper.units))}",
                f"Date: {paper.created_at.strftime('%Y-%m-%d')}",
            ],
        ]

        t = Table(data, colWidths=[3.5 * inch, 3 * inch])
        t.setStyle(
            TableStyle(
                [
                    ("GRID", (0, 0), (-1, -1), 1, colors.black),
                    ("PADDING", (0, 0), (-1, -1), 6),
                ]
            )
        )
        elements.append(t)
        elements.append(Spacer(1, 24))

        questions_data = paper.questions
        if not questions_data:
            elements.append(Paragraph("No questions generated.", normal_style))
            doc.build(elements)
            buffer.seek(0)
            return buffer.getvalue()

        mcqs = questions_data.get("mcqs", [])
        fill_blanks = questions_data.get("fill_blanks", [])
        short_qs = questions_data.get("short", [])
        long_qs = questions_data.get("long", [])

        if mcqs:
            elements.append(
                Paragraph("SECTION A: Multiple Choice Questions", header_style)
            )
            elements.append(Spacer(1, 12))

            for i, q in enumerate(mcqs, 1):
                q_text = f"{i}. {q['question']} ({q['marks']} marks)"
                elements.append(Paragraph(q_text, question_style))

                options = q["options"]
                if len(options) == 4:
                    opt_data = [
                        [f"a) {options[0]}", f"b) {options[1]}"],
                        [f"c) {options[2]}", f"d) {options[3]}"],
                    ]
                    opt_table = Table(opt_data, colWidths=[3 * inch, 3 * inch])
                    opt_table.setStyle(
                        TableStyle(
                            [
                                ("LEFTPADDING", (0, 0), (-1, -1), 20),
                                ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
                            ]
                        )
                    )
                    elements.append(opt_table)
                else:
                    for j, opt in enumerate(options):
                        label = chr(97 + j)
                        elements.append(Paragraph(f"{label}) {opt}", option_style))

                elements.append(Spacer(1, 12))

        if fill_blanks:
            elements.append(Paragraph("SECTION B: Fill in the Blanks", header_style))
            elements.append(Spacer(1, 12))

            for i, q in enumerate(fill_blanks, 1):
                q_text = f"{i}. {q['question']} ({q['marks']} marks)"
                elements.append(Paragraph(q_text, question_style))
                elements.append(Spacer(1, 6))

        if short_qs:
            elements.append(
                Paragraph("SECTION C: Short Answer Questions", header_style)
            )
            elements.append(Spacer(1, 12))

            for i, q in enumerate(short_qs, 1):
                q_text = f"{i}. {q['question']} ({q['marks']} marks)"
                elements.append(Paragraph(q_text, question_style))
                elements.append(Spacer(1, 24))

        if long_qs:
            elements.append(Paragraph("SECTION D: Long Answer Questions", header_style))
            elements.append(Spacer(1, 12))

            for i, q in enumerate(long_qs, 1):
                q_text = f"{i}. {q['question']} ({q['marks']} marks)"
                elements.append(Paragraph(q_text, question_style))
                elements.append(Spacer(1, 48))

        doc.build(elements)
        buffer.seek(0)
        return buffer.getvalue()

    @staticmethod
    def generate_docx(paper: PaperResponse, subject_name: str) -> bytes:
        buffer = io.BytesIO()
        doc = Document()

        heading = doc.add_heading("Question Paper", 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        table = doc.add_table(rows=2, cols=2)
        table.style = "Table Grid"

        row1 = table.rows[0].cells
        row1[0].text = f"Subject: {subject_name}"
        row1[1].text = f"Difficulty: {paper.difficulty.title()}"

        row2 = table.rows[1].cells
        row2[0].text = f"Units: {', '.join(map(str, paper.units))}"
        row2[1].text = f"Date: {paper.created_at.strftime('%Y-%m-%d')}"

        doc.add_paragraph()

        questions_data = paper.questions
        if not questions_data:
            doc.add_paragraph("No questions generated.")
            doc.save(buffer)
            buffer.seek(0)
            return buffer.getvalue()

        mcqs = questions_data.get("mcqs", [])
        fill_blanks = questions_data.get("fill_blanks", [])
        short_qs = questions_data.get("short", [])
        long_qs = questions_data.get("long", [])

        if mcqs:
            doc.add_heading("SECTION A: Multiple Choice Questions", level=1)
            for i, q in enumerate(mcqs, 1):
                p = doc.add_paragraph()
                p.add_run(f"{i}. {q['question']} ").bold = True
                p.add_run(f"({q['marks']} marks)")

                options = q["options"]
                opt_table = doc.add_table(rows=2, cols=2)

                if len(options) == 4:
                    opt_rows = opt_table.rows
                    opt_rows[0].cells[0].text = f"a) {options[0]}"
                    opt_rows[0].cells[1].text = f"b) {options[1]}"
                    opt_rows[1].cells[0].text = f"c) {options[2]}"
                    opt_rows[1].cells[1].text = f"d) {options[3]}"
                else:
                    for j, opt in enumerate(options):
                        label = chr(97 + j)
                        doc.add_paragraph(f"{label}) {opt}", style="List Bullet")

                doc.add_paragraph()

        if fill_blanks:
            doc.add_heading("SECTION B: Fill in the Blanks", level=1)
            for i, q in enumerate(fill_blanks, 1):
                p = doc.add_paragraph()
                p.add_run(f"{i}. {q['question']} ").bold = True
                p.add_run(f"({q['marks']} marks)")

        if short_qs:
            doc.add_heading("SECTION C: Short Answer Questions", level=1)
            for i, q in enumerate(short_qs, 1):
                p = doc.add_paragraph()
                p.add_run(f"{i}. {q['question']} ").bold = True
                p.add_run(f"({q['marks']} marks)")
                doc.add_paragraph()

        if long_qs:
            doc.add_heading("SECTION D: Long Answer Questions", level=1)
            for i, q in enumerate(long_qs, 1):
                p = doc.add_paragraph()
                p.add_run(f"{i}. {q['question']} ").bold = True
                p.add_run(f"({q['marks']} marks)")
                doc.add_paragraph()
                doc.add_paragraph()

        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
